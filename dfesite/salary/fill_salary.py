import os
import re
import docx
import dateparser
from datetime import timedelta, datetime
import pandas as pd
import numpy as np
from bs4 import BeautifulSoup

# from django.db import transaction
from django.conf import settings # correct way for access BASE_DIR, MEDIA_DIR...
# os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'dfesite.settings')
# import django
# django.setup()

from price.class_webnews import NewsStat, NewsStatDetail
from price.class_filehandle import WebFile, DocxFile
from industry import send_msg
from salary.models import SalaryNews, Salary, SalaryHead

import urllib3
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

MEDIA = settings.MEDIA_DIR

HEADER = {'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) \
          AppleWebKit/537.36 (KHTML, like Gecko) Chrome/62.0.3202.94 \
          Safari/537.36'}


# Функции добавления в БД
def add_news(title, news_href, date):
    SalaryNews.objects.create(title=title, href=news_href, pub_date=date)
    current_news = SalaryNews.objects.get(title=title)
    send_msg.sending('salary', current_news.id, current_news.title)
    return current_news.id


def add_head(ind, current_month_year, period_mm_year, pre_month, pre_year, pre_period, middle):
    SalaryHead.objects.create(salarynews_id=ind,
                                     current_my=current_month_year,
                                     period_mmy=period_mm_year,
                                     pre_month=pre_month,
                                     pre_year=pre_year,
                                     pre_period=pre_period,
                                     middle=middle)
    # return current


def add_data(ind, employer, current_zp, pre_month, pre_year, period, pre_period, middle):
    Salary.objects.create(salarynews_id=ind,
                                 employer=employer,
                                 current=current_zp,
                                 pre_month=pre_month,
                                 pre_year=pre_year,
                                 period=period,
                                 pre_period=pre_period,
                                 middle=middle)
    # return current


#---Проверка заголовков (сайт, файл)---
def check_header(docx_file, news_title):
    """ Ф-я проверки заголовка новости на сайте и в скачанном Word файле
    :return: 0 - не совпадают, 1 - совпадают
    """
    docx_header = ''
    for p, para in enumerate(docx_file.paragraphs):
        if p < 2:
            docx_header += para.text + ' '
    docx_header = re.sub('\d\)', '', del_s(docx_header))
    if docx_header == news_title:
        return 1
    return 0


# Вспомогательные функции: удаление символов, вычленение даты
def del_s(txt):
    """ 1. Удаление всех лишних пробелов, табуляций, новых строк
        2. Обработка '…', '-', коэффициента
        3. Преобразование текста в вещ.число
    """
    reg_ex = re.compile('в\s*\d{1,2}[,.]\d{1,2}\s*р')
    new_txt = re.sub('\s+', ' ', txt).strip()
    if new_txt == '':
        return float('nan')
    elif re.search('…', new_txt) or re.search('\.\.\.', new_txt):
        return 0.0123456789
    elif new_txt == '-':
        return 0.0123454321
    elif re.search(reg_ex, new_txt):  # если указано "в 1,6р", вычленяем 1,6
        koef = re.search('\d{1,2}[,.]\d{1,2}', new_txt).group()
        new_txt = 100 * float(re.sub(',', '.', koef))
        return round(new_txt, 1)

    try:
        new_txt = float(re.sub(',', '.', new_txt))
    except ValueError:
        pass

    return new_txt


def cut_date(txt):
    """
    Из заголовка вычленяем дату. Например из текста (txt):
        "...по видам экономической деятельности за январь-сентябрь 2020 года..."
    получим "сентябрь 2020". Ф-я возвращает эту строку в формате дата.
    """
    regex = re.compile('[яфмаисонд][а-я]+[ьтй]\s+\d{4}')
    try:
        dateparser.parse(re.search(regex, txt).group())
    except AttributeError:  # ...деятельности за 2018... преобразуется в (2018, 12, ...)
        regex = re.compile('\d{4}')
    return dateparser.parse(re.search(regex, txt).group())


#----------------------------------------------------
# Функции создания pandas dataframe из таблицы html
def pre_process_table(table):
    """
    INPUT:  1. table - a bs4 element that contains the desired table: ie <table> ... </table>
    OUTPUT: a tuple of:
            1. rows - a list of table rows ie: list of <tr>...</tr> elements
            2. num_rows - number of rows in the table
            3. num_cols - number of columns in the table
    Options: include_td_head_count - whether to use only th or th and td to count number of columns (default: False)
    """
    rows = [x for x in table.findAll('tr')]

    num_rows = len(rows)

    # get an initial column count. Most often, this will be accurate
    num_cols = max([len(x.findAll(['th','td'])) for x in rows])

    # sometimes, the tables also contain multi-colspan headers. This accounts for that:
    header_rows_set = [x.findAll(['th', 'td']) for x in rows if len(x.findAll(['th', 'td']))>num_cols/2]

    num_cols_set = []

    for header_rows in header_rows_set:
        num_cols = 0
        for cell in header_rows:
            row_span, col_span = get_spans(cell)
            num_cols+=len([cell.getText()]*col_span)

        num_cols_set.append(num_cols)

    num_cols = max(num_cols_set)

    return rows, num_rows, num_cols


def get_spans(cell):
        """
        INPUT:  1. cell - a <td>...</td> or <th>...</th> element that contains a table cell entry
        OUTPUT: 1. a tuple with the cell's row and col spans
        """
        if cell.has_attr('rowspan'):
            rep_row = int(cell.attrs['rowspan'])
        else: # ~cell.has_attr('rowspan'):
            rep_row = 1
        if cell.has_attr('colspan'):
            rep_col = int(cell.attrs['colspan'])
        else: # ~cell.has_attr('colspan'):
            rep_col = 1

        return rep_row, rep_col


def process_rows(rows, num_rows, num_cols):
    """
    INPUT:  1. rows - a list of table rows ie <tr>...</tr> elements
    OUTPUT: 1. data - a Pandas dataframe with the html data in it
    """
    data = pd.DataFrame(np.ones((num_rows, num_cols))*np.nan)
    for i, row in enumerate(rows):
        try:
            col_stat = data.iloc[i,:][data.iloc[i,:].isnull()].index[0]
        except IndexError:
            print(i, row)

        for j, cell in enumerate(row.findAll(['td', 'th'])):
            rep_row, rep_col = get_spans(cell)

            #find first non-na col and fill that one
            while any(data.iloc[i, col_stat:col_stat+rep_col].notnull()):
                col_stat += 1

            data.iloc[i:i+rep_row, col_stat:col_stat+rep_col] = cell.getText()
            if col_stat < data.shape[1]-1:
                col_stat += rep_col

    return data

#----------------------------------------------------
# Функции создания pandas dataframe из таблицы docx
def read_docx_table(document):
    table = document.tables[0]
    data = [[cell.text for cell in row.cells] for row in table.rows]
    if len(table.columns) > 5:
        df = pd.DataFrame(data, columns=[0, 1, 2, 3, 4, 5, 6])
    else:
        df = pd.DataFrame(data, columns=[0, 1, 2, 3, 4])
    return df


#----------------------------------------------------
def db_addhead(news_pk, max_cols, df):
    if max_cols > 5:
        for row, col in df.iterrows():
            if row == 0:
                hcur_my = del_s(col[1])
                hper_mmy = del_s(col[4])
            elif row == 1:
                pass
            elif row == 2:
                hpre_m = del_s(col[2])
                hpre_y = del_s(col[3])
                hper_m = del_s(col[5])
                hmiddle = del_s(col[6])
                add_head(news_pk, hcur_my, hper_mmy, hpre_m, hpre_y, hper_m, hmiddle)
                # print(row, hcur_my, hper_mmy, hpre_m, hpre_y, hper_m, hmiddle)
            else:
                break
    else:
        for row, col in df.iterrows():
            if row == 0:
                hcur_my = del_s(col[1])
                hper_mmy = del_s(col[2])
            elif row == 1:
                hpre_m = del_s(col[2])
                hpre_y = del_s(col[3])
                hmiddle = del_s(col[4])
                add_head(news_pk, hcur_my, hper_mmy, hpre_m, hpre_y, 'NaN', hmiddle)
                # print(row, hcur_my, hper_mmy, hpre_m, hpre_y, 'NaN', hmiddle)
            else:
                break


def db_adddocx(news_pk, max_cols, df):
    nan = float('nan')
    if max_cols > 5:
        for row, col in df.iterrows():
            if row > 2:
                add_data(news_pk, del_s(col[0]), del_s(col[1]), del_s(col[2]), del_s(col[3]), del_s(col[4]), del_s(col[5]), del_s(col[6]))
    else:
        for row, col in df.iterrows():
            if row > 1:
                add_data(news_pk, del_s(col[0]), del_s(col[1]), del_s(col[2]), del_s(col[3]), nan, nan, del_s(col[4]))


# 2020-2018. Загрузка данных из ранее скачанных файлов
def db_addhtml(news_pk, max_cols, df):
    """ Добавление заголовка и данных таблицы html в БД """
    nan = float('nan')
    rbegin = 300
    rend = 300
    nao_begin = re.compile("Ненецкий")
    nao_end = re.compile("Ненецкого")
    if max_cols > 5:  # колонок 7
        for row, col in df.iterrows():
            if row > 10 and re.search(nao_begin, col[0]):
                rbegin = row
                add_data(news_pk, del_s(col[0]), del_s(col[1]), del_s(col[2]), del_s(col[3]), del_s(col[4]), del_s(col[5]), del_s(col[6]))
                # print(row, del_s(col[0]), del_s(col[1]), del_s(col[2]), del_s(col[3]), del_s(col[4]), del_s(col[5]), del_s(col[6]))
            elif re.search(nao_end, col[0]):
                rend = row
            elif rbegin <= row < rend:
                add_data(news_pk, del_s(col[0]), del_s(col[1]), del_s(col[2]), del_s(col[3]), del_s(col[4]), del_s(col[5]), del_s(col[6]))
                # print(row, del_s(col[0]), del_s(col[1]), del_s(col[2]), del_s(col[3]), del_s(col[4]), del_s(col[5]), del_s(col[6]))
            elif rbegin > rend:
                add_data(news_pk, del_s(col[0]), del_s(col[1]), del_s(col[2]), del_s(col[3]), del_s(col[4]), del_s(col[5]), del_s(col[6]))
                # print(row, del_s(col[0]), del_s(col[1]), del_s(col[2]), del_s(col[3]), del_s(col[4]), del_s(col[5]), del_s(col[6]))
    else:  # колонок 5
        for row, col in df.iterrows():
            if row > 10 and re.search(nao_begin, col[0]):
                rbegin = row
                add_data(news_pk, del_s(col[0]), del_s(col[1]), del_s(col[2]), del_s(col[3]), nan, nan, del_s(col[4]))
                # print(row, del_s(col[0]), del_s(col[1]), del_s(col[2]), del_s(col[3]), nan, nan, del_s(col[4]))
            elif re.search(nao_end, col[0]):
                rend = row
            elif rbegin <= row < rend:
                add_data(news_pk, del_s(col[0]), del_s(col[1]), del_s(col[2]), del_s(col[3]), nan, nan, del_s(col[4]))
                # print(row, del_s(col[0]), del_s(col[1]), del_s(col[2]), del_s(col[3]), nan, nan, del_s(col[4]))
            elif rbegin > rend:
                add_data(news_pk, del_s(col[0]), del_s(col[1]), del_s(col[2]), del_s(col[3]), nan, nan, del_s(col[4]))
                # print(row, del_s(col[0]), del_s(col[1]), del_s(col[2]), del_s(col[3]), nan, nan, del_s(col[4]))


def from_htmldocx(path):
    # news_id = 1
    for i, filename in enumerate(os.listdir(path)):
        file_path = os.path.join(path, filename)
        if filename.endswith(".htm"):
            soup = BeautifulSoup(open(file_path), 'html.parser')
            table = soup.find('table')
            rows, num_rows, num_cols = pre_process_table(table)
            dataframe = process_rows(rows, num_rows, num_cols)
            title = del_s(soup.title.text)
            pub_date = (cut_date(title).replace(day=1) + timedelta(days=32)).replace(day=1)
            news_id = add_news(title, '', pub_date)
            db_addhead(news_id, len(dataframe.columns), dataframe)
            print(f"news_id={news_id}, заголовок добавлен")
            db_addhtml(news_id, num_cols, dataframe)
        elif filename.endswith(".docx"):
            docx_header = ''
            doc = docx.Document(file_path)
            for p, para in enumerate(doc.paragraphs):
                if p < 2:
                    docx_header += para.text + ' '
            title = re.sub('\d\)', '', del_s(docx_header))
            pub_date = (cut_date(title).replace(day=1) + timedelta(days=32)).replace(day=1)
            dataframe = read_docx_table(doc)
            news_id = add_news(title, '', pub_date)
            db_addhead(news_id, len(dataframe.columns), dataframe)
            print(f"news_id={news_id}, заголовок добавлен")
            db_adddocx(news_id, len(dataframe.columns), dataframe)


#----------------------------------------------------
# Добавление данных из файлов htlm (2018 - 2019_07), docx (2019_07 - 2020_07)
# files_dir = "d:/code/python/django/dfeproject/dfesite/media/salary/_source"
# from_htmldocx(files_dir)
# print('Процедура завершена')
#----------------------------------------------------

# Добавление данных с сайта. Поиск последней добавленной новости
def last_added_news(news_text):
    page = 1
    try:
        last_db_date = SalaryNews.objects.last().pub_date
    except AttributeError:
        last_db_date = datetime(2019, 9, 16, 0, 0)
    while True:
        url = 'https://arhangelskstat.gks.ru/news?page=' + str(page)
        stat_news = NewsStat(0, news_text, url, HEADER)
        if stat_news.atag is not None:
            news_date = stat_news.get_pub_date()
            print(f'page={page}, site={news_date}, db={last_db_date}')
            print(f'{stat_news.atag}')
            if news_date == last_db_date:
                return page - 1
        page += 1


# ---Добавление данных с сайта. Поиск новости, обработка---
def search_news(page, news_text):
    """
    Поиск новости news_text
    :param page: номер страницы с новостями статистики
    :param news_text: искомая новость
    :return: list[news_count, title, href, pub_date, file]
    """
    nao = 'Ненецком'
    app_dir = 'salary'
    webpage = 'https://arhangelskstat.gks.ru/news?page=' + str(page)
    news = []
    # 0-количество новостей, 1-заголовок, 2-ссылка, 3-дата, 4-файл (либо путь к xl, либо объект docx)
    stat = NewsStat(0, news_text, webpage, HEADER)  # аргумент 0 - это 1ая новость
    if stat.acount == 0:
        return None
    news.append(stat.acount)
    news.append(stat.get_title())
    news.append(stat.get_href())
    news.append(stat.get_pub_date())
    stat_detail = NewsStatDetail(nao, stat.get_href(), HEADER)
    year = re.findall(r'\d{4}', stat.get_title())[0]
    stat_file = WebFile(stat_detail.file_href, MEDIA, app_dir, year, stat_detail.file_name)
    stat_file.download_file()
    file = DocxFile(MEDIA, app_dir, year, stat_detail.file_name).get_docx()
    news.append(file)
    return news


def populate():
    news_find = 'реднемесячная номинальная начисленная заработная плата'
    page_num = last_added_news(news_find)
    while page_num > 0:
        # 0-количество новостей, 1-заголовок, 2-ссылка, 3-дата, 4-файл (либо путь к xl, либо объект docx)
        newsdata = search_news(page_num, news_find)
        if not newsdata:
            page_num -= 1
            continue
        news_id = add_news(newsdata[1], newsdata[2], newsdata[3])

        test_head = check_header(newsdata[4], newsdata[1])
        if not test_head:
            print("ВНИМАНИЕ! Проверьте заголовки новости на сайте и в скачанном в файле. \n"
                  "Возможно, изменилась структура файла.")

        dataframe = read_docx_table(newsdata[4])
        db_addhead(news_id, len(dataframe.columns), dataframe)
        db_adddocx(news_id, len(dataframe.columns), dataframe)

        page_num -= 1
    print('Процедура выполнена')


# if __name__ == "__main__":
#     populate()
