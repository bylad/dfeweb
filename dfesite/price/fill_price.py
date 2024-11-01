import os
import sys
import re
import xlrd
import docx
import openpyxl
import pandas as pd
import dateparser
from datetime import timedelta

from django.db import transaction
from django.conf import settings # correct way for access BASE_DIR, MEDIA_DIR...

from .class_webnews import NewsStat, NewsStatDetail, PriceStat
from .class_filehandle import WebFile, DocxFile
from .models import PriceNews, PriceData, PricePetrolHead, PricePetrolData
from industry import send_msg
from dfesite.constants import HEADER, MONTHS

import urllib3
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

MEDIA = settings.MEDIA_DIR
SEARCH_TXT1 = 'Ненецкий автономный округ'
SEARCH_TXT2 = 'Нарьян-Мар'
CUT_NM = ' и г.Нарьян-Мару'
SEARCH_NEWS_TXT = ' на отдельные потребительские товары'
SEARCH_PETROL_TXT = ' на бензин'
WEBPAGE = 'https://29.rosstat.gov.ru/prices111'


# Функции добавления в БД
def add_news(title, news_href, date):
    PriceNews.objects.get_or_create(title=title, href=news_href, pub_date=date)
    current_news = PriceNews.objects.get(title=title)
    #send_msg.sending('price', current_news.id, current_news.title)  # здесь срабатывала ложно
    return current_news.id


def add_data(ind, tovar, sale_price):
    PriceData.objects.get_or_create(pricenews_id=ind, product=tovar, price=sale_price)
    return tovar


def add_pethead(ptitle, phref, pdate):
    PricePetrolHead.objects.get_or_create(petrol_title=ptitle, href=phref, pub_date=pdate)
    current_pethead = PricePetrolHead.objects.get(petrol_title=ptitle)
    return current_pethead.id


def add_petdata(ind, toplivo, sale_price):
    PricePetrolData.objects.get_or_create(pricepetrolhead_id=ind, petrol=toplivo, price=sale_price)
# ------------------------


def cut_date(txt):
    """
    Из заголовка вычленяем дату. Например из текста (txt):
        "Средние  цены на отдельные потребительские товары (услуги)
         по Ненецкому автономному округу на 13 января 2020 года"
    получим "13 января 2020". Ф-я возвращает эту строку в формате дата.
    """
    regex = re.compile('\d{1,2}\s[яфмаисонд][а-я]+[а|я]\s\d{4}')
    return dateparser.parse(re.search(regex, txt).group())


# Функции обработки файлов DOCX, XLS(X)
def data_docx(doc):
    pet_price = [0.0] * 3
    pet_name = [''] * 3
    ben_i = 1000
    j = 0
    ben = ['Бензин автомобильный марки АИ-92']
    if len(doc.tables) > 1:
        print('Внимание! Изменилось количество таблиц. Проверьте исходный файл.')
        print('Обрабатывается 2-ая таблица...')
    for i, row in enumerate(doc.tables[1].rows):
        for cell in row.cells:  # поиск строки с необх. данными
            if any(x in cell.text for x in ben):
                ben_i = i
                break
        if ben_i <= i < len(doc.tables[1].rows) and j < 3:
            pet_name[j] = doc.tables[1].cell(i, 0).text
            pet_price[j] = float(re.sub(',', '.', doc.tables[1].cell(i, 1).text))
            j += 1
    pet_head = doc.tables[0].cell(0, 0).text.replace(CUT_NM, '')
    print('Обработка таблицы DOCX файла успешно завершена.\n')
    return pet_head, pet_name, pet_price


def read_xl(xlsx):
    df = pd.read_excel(xlsx)
    if any(df.iloc[:, 0].str.match(SEARCH_TXT1)):
        xls_title = re.sub(r'\s+', ' ', df.columns[0].replace(CUT_NM, ''))
        df.iloc[:, 0] = df.iloc[:, 0].str.strip()  # удаляем начальные и конечные
        row_first = df.index[df.iloc[:, 0].str.match(SEARCH_TXT1)][0]
        row_last = df.index[df.iloc[:, 0].str.contains(SEARCH_TXT2)][0]
        df1 = df.iloc[row_first + 1:row_last, 0:2]
        print('Обработка Excel файла успешно завершена.\n')
        return xls_title, df1[df1.columns[0]].tolist(), df1[df1.columns[1]].tolist()
    print()
    return "", None, None


def data_openpyxl(xlsx, search_txt):
    j = 0
    data_i = 0
    petrol_list = ['Бензин автомобильный марки АИ-92', 'Бензин автомобильный марки АИ-95', 'Дизельное топливо']
    regex = re.compile(search_txt)
    wb = openpyxl.load_workbook(xlsx)
    ws = wb.worksheets[0]
    xls_title = re.sub('\s+', ' ', ws.cell(row=1, column=1).value)  # удаляем все лишние пробелы, символы новой строки
    for i in range(1, ws.max_row):
        if re.search(regex, ws.cell(row=i, column=1).value):
            data_i = i
            print(f'data_i={data_i}')
            break
    if not data_i:
        print('Ошибка!\nНе найден заголовок "Товар". Проверьте файл!\nПрограмма завершена.')
        sys.exit()

    rows_count = ws.max_row - data_i
    col_name = [''] * rows_count
    col_price = [0.0] * rows_count

    # NEW ADD BEGIN
    if search_txt == 'Нарьян-Мар':
        col_name = [''] * 3
        col_price = [0.0] * 3

        for i in range(data_i+1, ws.max_row):  # итерируем цены для Нарьян-Мара
            if any(ext in ws.cell(row=i, column=1).value for ext in petrol_list):
                idx = [petrol_list.index(s) for s in petrol_list if s in ws.cell(row=i, column=1).value]
                col_name[idx[0]] = ws.cell(row=i, column=1).value
                col_price[idx[0]] = ws.cell(row=i, column=4).value
        print('Обработка Excel файла успешно завершена.\n')
        return col_name, col_price
    # NEW ADD END

    for i in range(data_i+1, ws.max_row):
        col_name[j] = ws.cell_value(i, 0)
        col_price[j] = ws.cell_value(i, 1)
        j += 1

    print('Обработка Excel файла успешно завершена.\n')
    return xls_title, col_name, col_price
#--------------------------------


def check_db(news_date, is_petrol):
    """ Проверка наличия данных в базе (0 - нет, 1 - есть)
    is_petrol: цены на бензин (0 - нет, 1 - да)
    """
    data_exists = 1
    if is_petrol:
        db_data = PricePetrolHead.objects.all()
    else:
        db_data = PriceNews.objects.all()
    for i in range(len(db_data)-1, 0, -1):  # начинаем проверку с даты, добавленной последним
        if news_date.date() == db_data[i].pub_date.date():
            data_exists = 1
            print(f'Данные уже в базе!')
            break
        else:
            data_exists = 0
    return data_exists


# ---Добавление данных с сайта---
def search_price(idx, page, news_text):
    """
    Поиск новости news_text
    :param idx: номер новости на веб-странице
    :param page: номер страницы с новостями статистики
    :param news_text: искомая новость
    :return: list[news_count, title, href, pub_date, file]
    """
    app_dir = 'price'
    price_avg = []
    # 0-количество новостей, 1-заголовок, 2-ссылка, 3-дата, 4-файл (либо путь к xl, либо объект docx)
    stat = PriceStat(idx, news_text, page, HEADER)
    if stat.div_count:
        price_avg.append(stat.div_count)
        price_avg.append(stat.get_title())
        if stat.get_href() == "#":  # Если нет ссылки, то оставляем общую
            price_avg.append(stat.webfile_link)
        else:
            price_avg.append(stat.get_href())
        price_avg.append(stat.get_pub_date())
        year = re.findall(r'\d{4}', stat.get_title())[0]
        stat_file = WebFile(stat.webfile_href, MEDIA, app_dir, year, stat.webfile_name)
        stat_file.download_file()
        file = stat_file.file_path
        file_extension = os.path.splitext(file)[1]
        if news_text == SEARCH_PETROL_TXT:
            if file_extension.lower() == '.docx' or file_extension.lower() == '.doc':
                file = DocxFile(MEDIA, app_dir, year, os.path.split(file)[1]).get_docx()
        price_avg.append(file)

        return price_avg
    return None


def mid_data(news_num):
    """
    Поиск статданных 'Средние цены на отдельные потребительские товары'
    с индексом (news_num) на веб странице 'https://29.rosstat.gov.ru/news?page=' + str(page)
    Добавление найденных данных в БД
    Возвращает количество найденных новостей на странице
    0-количество новостей, 1-заголовок, 2-ссылка, 3-дата, 4-файл (либо путь к xl, либо объект docx)
    """
    news_id = None
    newsdata = search_price(news_num, WEBPAGE, SEARCH_NEWS_TXT)
    # all_news = PriceNews.objects.all().order_by('-pub_date')  # .all() - следует избегать такой конструкции
    lastdb_news = PriceNews.objects.last()
    all_news = PriceNews.objects.filter(pub_date__year=f"{lastdb_news.pub_date.year}").order_by('-pub_date')
    if newsdata is not None:
        for news in all_news:  # добавлена проверка, т.к. на сайте статистики м.б. ошибочная дата в заголовке
            if news.title == newsdata[1] and news.pub_date != newsdata[3]:
                previous_news = news.get_previous_by_pub_date()
                pattern = re.compile(r'\d{1,2}\s[яфмаисонд][а-я]+[а|я]\s\d{4}')
                previous_news_title_date = dateparser.parse(re.search(pattern, previous_news.title).group())
                # т.к. данные обновляются еженедельно, то к предыдущей дате добавляем 7 дней
                newsdate = previous_news_title_date + timedelta(days=7)
                new_date = f"{newsdate.day - 2} {MONTHS[newsdate.month - 1]} {newsdate.year}"
                newsdata[1] = re.sub(pattern, new_date, newsdata[1])
            elif news.title == newsdata[1] and news.pub_date == newsdata[3]:
                data_in = 1
                return newsdata[0], data_in, 0  # news_id=0, т.к. новость уже существует

        data_in = check_db(newsdata[3], 0)
        if data_in == 0:
            # определяем ID новости, к нему будут привязаны данные
            news_id = add_news(newsdata[1], newsdata[2], newsdata[3])
            xls_file = newsdata[4]
            xl_title, products, prices = read_xl(xls_file)
            for p in range(len(products)):  # кол-во строк с ценами на товары
                add_data(news_id, products[p], prices[p])
        return newsdata[0], data_in, news_id

    return 0, None, None


def pet_data(news_num):
    """
    Поиск новости 'О потребительских ценах на бензин' на странице 'https://29.rosstat.gov.ru/prices111'
    Добавление найденных данных в БД
    Возвращает количество найденных новостей на странице
    """
    # [0] количество новостей, [1] заголовок, [2] ссылка, [3] дата, [4] файл
    newsdata = search_price(news_num, WEBPAGE, SEARCH_PETROL_TXT)
    if newsdata is not None:
        data_in = check_db(newsdata[3], 1)
        if data_in == 0:
            pethead_id = add_pethead(newsdata[1], newsdata[2], newsdata[3])
            if not isinstance(newsdata[4], str):
                print('============== pet_news: DOC FILE ===============')
                pet_title, pet_name, pet_price = data_docx(newsdata[4])

                for p in range(len(pet_name)):  # кол-во строк с ценами на товары
                    add_petdata(pethead_id, pet_name[p], pet_price[p])
            else:
                print('============== pet_news: XLSX FILE ===============')
                pet_name, pet_price = data_openpyxl(newsdata[4], 'Нарьян-Мар')
                for p in range(len(pet_name)):  # кол-во строк с ценами на товары
                    add_petdata(pethead_id, pet_name[p], pet_price[p])
        return newsdata[0]
    return 0


def from_web():
    mid_num = 0
    mid_count = 1
    pet_num = 0
    pet_count = 1
    data_in_db = 0

    while mid_num < mid_count:
        print(f'mid_data RUNNING... mid_num={mid_num} < mid_count={mid_count}')
        mid_count, data_in_db, midnews_id = mid_data(mid_num)
        mid_num += 1

    while pet_num < pet_count:
        print(f'pet_news RUNNING... pet_num={pet_num} < pet_count={pet_count}')
        pet_count = pet_data(pet_num)
        pet_num += 1
    return data_in_db, midnews_id


@transaction.atomic
def populate():
    print('============== PRICE BEGIN ===============')
    try:
        data_exist, mnews_id = from_web()
        if data_exist == 1:
            print('===========PRICE DATA EXIST===========')
        current_news = PriceNews.objects.get(id=mnews_id)
        if current_news:
            print(f'current_news.id={current_news.id}')
            send_msg.sending('price', current_news.id, current_news.title)
    except Exception as e:
        print("fillprice.populate() procedure Exception error:")
        print(e)
    print('============== PRICE END ===============')


#Data adding into database from files
#def mid_data(wrkdir):
#    """ Добавление средних цен товаров из xlsx """
#    del_naonm = "по Ненецкому автономному округу и г. Нарьян-Мару "
#    goods_files = [f for f in os.listdir(wrkdir) if ".xlsx" in f.lower()]
#    print("FIIIIIIIIIIIIIIIIIIIIIIIIIIIIILLLLLLLLLLLLLLLLLLLLLLLLLLLEEEEEEEEEEEEEEEEEEE")
#    print(goods_files)
#    goods_all = PriceNews.objects.all().order_by("pub_date")
#    print(goods_all)
#
#    for f in goods_files:
#        xl_title, products, prices = read_xl(f"{wrkdir}/{f}")
#        goods_title = xl_title.replace(del_naonm, "")
#        for good in goods_all:
#            if goods_title == good.title:
#                for p in range(len(products)):  # кол-во строк с ценами на товары
#                    add_data(good.id, products[p], prices[p])
#        goods_pubdate = cut_date(goods_title) + timedelta(days=2)  # публикуется на 2 дня позже
#        news_id = add_news(goods_title, WEBPAGE, goods_pubdate)
#        for p in range(len(products)):  # кол-во строк с ценами на товары
#            add_data(news_id, products[p], prices[p])
#
#
#def pet_data(wrkdir):
#    """ Добавление средних цен на топливо из docx """
#    from docx import Document
#    del_str = [" по Ненецкому автономному округу и городу Нарьян-Мару", " (по выборочному кругу автозаправочных станций)"]
#    petrol_files = [f for f in os.listdir(wrkdir) if ".docx" in f.lower()[-5:]]  # последние 5 символов
#    for f in petrol_files:
#        pet_title, pet_name, pet_price = data_docx(Document(f"{wrkdir}/{f}"))
#        pet_title = re.sub(r'\s+', ' ', pet_title).strip()  # нормализация заголовка
#        petrol_title = pet_title.replace(del_str[0], "").replace(del_str[1], "")  # удаление из заголовка
#        petrol_pubdate = cut_date(petrol_title) + timedelta(days=2)     # дата публикации на 2 дня позже
#        pethead_id = add_pethead(petrol_title, WEBPAGE, petrol_pubdate)
#        for p in range(len(pet_name)):  # кол-во строк с ценами на товары
#            add_petdata(pethead_id, pet_name[p], pet_price[p])
#
#
#@transaction.atomic
#def populate():
#    """ Data adding into database from files """
#    print("-----------------------PRICE BEGIN--------------------------")
#    work_dir = "/home/sa/projects/dfeweb/dfesite/media/price/2024/lost"
#    mid_data(work_dir)
#    pet_data(work_dir)
#    print("-----------------------PRICE END--------------------------")
