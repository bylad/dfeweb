import os
import re, requests
import subprocess
from bs4 import BeautifulSoup
import dateparser
import docx
from datetime import datetime as dt
from transliterate import translit

from django.conf import settings # correct way for access BASE_DIR, MEDIA_DIR...
from dfesite.constants import HEADER, MONTH, MONTHE
from .models import (IndustryNews, IndustryIndex, IndustryProduction, IndustryIndexHead, IndustryProductionHead)
from . import  send_msg
from django.db import transaction

MEDIA = settings.MEDIA_DIR
WEBPAGE = 'https://29.rosstat.gov.ru/production111'  # Поиск оперативной инфоромации


class NewsLocate:
    def __init__(self, web, HEADER, txt):
        self.request = requests.get(web, headers=HEADER)
        self.soup = BeautifulSoup(self.request.content, 'html.parser')
        self.atag = self.soup.find('a', text=re.compile(txt))
        if self.atag is not None:
            atag_parent = self.atag.parent.parent
            try:
                news_date = translit(atag_parent.find('div', class_='news-card__data').text,'ru')
            except AttributeError:
                news_date = '9 сентября 1999'
                print("Attribute Error! News date was dropped!")
            self.publicated = dateparser.parse(news_date)


def date_int(newstitle):
    """ Вычленяем из заголовка новости год и порядковые номера месяцев
        dig0-год, dig1-месяц1, dig2-месяц2
    """
    dig = [int(s) for s in newstitle.split() if s.isdigit()]
    for string in MONTHE:
        regex = re.compile(string)
        match = re.search(regex, newstitle)
        if match:
            mesyac = newstitle[match.start():match.end()]
            monthnum = '{:02}'.format(MONTHE.index(mesyac) + 1)
            dig.append(int(monthnum))
    return dig


def doc2docx(basedir):
    """ Преобразуем файл DOC в DOCX
    """
    for dir_path, dirs, files in os.walk(basedir):
        for file_name in files:
            file_path = os.path.join(dir_path, file_name)
            file_name, file_extension = os.path.splitext(file_path)
            if file_extension.lower() == '.doc':
                docx_file = '{0}{1}'.format(file_path, 'x')
                # Skip conversion where docx file already exists
                if not os.path.isfile(docx_file):
                    print(f'Преобразование в docx\n{file_path}\n')
                    try:
                        os.chdir(basedir)
                        subprocess.call(['lowriter', '--convert-to', 'docx', file_path])
                    except Exception:
                        print('Failed to Convert: {file_path}')


def table_doc(doc):
    data = [[],[]]
    for t in range(2):  # ранее использовали len(doc.tables), отказался, т.к. в подписи могут добавить еще таблицу
        table = doc.tables[t]
        table_columns = len(table.columns)
        # В янв 2021г. 3 столбца, по факту 4 (объединены). В янв 2022 объединение удалено.
        if t == 0:
            if table_columns != 3 or table_columns != 4: 
                print(f'Внимание! Изменилась структура таблицы, проверьте скачанный файл.')
                print(f'Количество столбцов в таблице t[{t}] = {table_columns}.')
            # os.system("pause")
        elif t == 1 and table_columns != 3:
            print(f'Внимание! Изменилась структура таблицы, проверьте скачанный файл.')
            print(f'Количество столбцов в таблице t[{t}] = {table_columns}.')
            # os.system("pause")

        for i, row in enumerate(table.rows):
            text = (cell.text for cell in row.cells)
            row_data = list(text)
            data[t].append(row_data)
    return data[0], data[1]


def floating(start, t):
    """
    В значении оставляем только цифры и запятую, затем меняем ',' на '.'
    Перевод из str в float. При отсутствии значения '...' задать '0,0'
    :param start: индекс начала итерации (табл.1 = 2, табл.2 = 1), пропуск заголовка
    :param t: таблица в виде списка
    """
    for row in range(start, len(t)):
        for col in range(1, len(t[1])): # обработка со второго столбца
            if t[row][col][:3] == '...' or t[row][col].strip() == '-' or t[row][col].strip() == '':
                t[row][col] = '0,0'
            t[row][col] = float(re.sub('[^0-9,]', "", t[row][col]).replace(",", "."))
    return t


def add_news(data_title, data_href, data_date):
    IndustryNews.objects.get_or_create(title=data_title, href=data_href, pub_date=data_date)
    current_news = IndustryNews.objects.get(title=data_title)
    send_msg.sending('industry', current_news.id, current_news.title)
    return current_news.id


def add_table_index_HEADER(ind, moye, pre, cur, precur):
    IndustryIndexHead.objects.get_or_create(industrynews_id=ind,
                                            month_year=moye,
                                            pre_year=pre,
                                            cur_year=cur,
                                            pre_cur=precur)


def add_table_production_HEADER(ind, cur, precur):
    IndustryProductionHead.objects.get_or_create(industrynews_id=ind,
                                                 cur_year=cur,
                                                 pre_cur=precur)


def add_table_index(ind, name, pre, cur, precur):
    IndustryIndex.objects.get_or_create(industrynews_id=ind,
                                        production_index=name,
                                        pre_year_index=pre,
                                        cur_year_index=cur,
                                        pre_cur_index=precur)
    return name


def add_table_production(ind, name, cur, precur):
    IndustryProduction.objects.get_or_create(industrynews_id=ind,
                                             industry_production=name,
                                             cur_year_production=cur,
                                             pre_cur_production=precur)
    return name


def create_db(id_news, t1, t2, is_january):
    if is_january:  # В таблице с индексом для января используем 3 столбца, одну зануляем
        # Создание заголовков таблиц
        if t1[0][1] and t1[1][1] and t1[1][2]:
            add_table_index_HEADER(id_news, t1[0][1], t1[1][1], t1[1][2], 0)
        else:
            print(f'Внимание! Изменилась структура таблицы с индексом, проверьте ее заголовки.')
        # Создание в БД таблицы с индексом
        t1_columns_len = len(t1[2])
        if t1_columns_len == 3:
            for row in range(2, len(t1)):  # в янв.2021 3 столбца
                add_table_index(id_news, t1[row][0], 0, t1[row][1], t1[row][2])
        else:  # если ошибочно в январе используется 4 столбца вместо 3
            for row in range(2, len(t1)):  # зануляем 2ой столбец, т.к. 2 и 3 повторяются
                add_table_index(id_news, t1[row][0], 0, t1[row][2], t1[row][3])
    else:  # Создание заголовков таблиц для остальных месяцев, используем 4 столбца
        if t1[0][1] and t1[1][1] and t1[1][2] and t1[1][3]:
            add_table_index_HEADER(id_news, t1[0][1], t1[1][1], t1[1][2], t1[1][3])
        else:
            print(f'Внимание! Изменилась структура таблицы с индексом, проверьте ее заголовки.')
        # Заполняем в БД тело таблицы с индексом
        for row in range(2, len(t1)):
            add_table_index(id_news, t1[row][0], t1[row][1], t1[row][2], t1[row][3])
    # Заголовок и тело таблицы с производством
    if t2[0][1] and t2[0][2]:
        add_table_production_HEADER(id_news, t2[0][1], t2[0][2])
    else:
        print(f'Внимание! Изменилась структура таблицы с производством, проверьте ее заголовки.')
    for row in range(1, len(t2)):
        add_table_production(id_news, t2[row][0], t2[row][1], t2[row][2])


# Функция добавления данных по производству
def last_added_news(header):
    dict_key = None
    try:
        last_db_date = IndustryNews.objects.last().pub_date
    except AttributeError:  # datetime.now().year - текущий год в формате int
        last_db_date = dt(dt.now().year - 1, 1, 1, 0, 0)

    url = WEBPAGE
    stat_news = NewsLocate(url, header, 'О промышленном производстве')

    # Создание списка тегов div из найденных заголовков
    all_inds = stat_news.soup.find_all('div', text=re.compile('О промышленном производстве'))

    # re.sub(r'\s+', '', x.text).strip(): x.text - выделение заголовка. Удаление перевода строк, нач. и конечн. пробелов
    industry_heads = [re.sub(r'\s+', ' ', x.text).strip() for x in all_inds if str(dt.now().year) in x.text]

    # WEBPAGE.rsplit('/', 1)[0] - разделение строки на 2 подстроки, где символом разделения является слеш справа
    # x.find_parent().find_parent().find('a').get('href') - переход на 2 позиции выше, поиск тега 'a', извлечение 'href'
    industry_hrefs = [f"{WEBPAGE.rsplit('/', 1)[0]}{x.find_parent().find_parent().find('a').get('href')}"
                      for x in all_inds if str(dt.now().year) in x.text]

    date_pattern = r'\d{2}\.\d{2}\.\d{4}'
    # Извлечение даты как текст и преобразование в datetime
    industry_dates = [dt.strptime(re.search(date_pattern, x.find_next_sibling().text).group(0), '%d.%m.%Y')
                      for x in all_inds if str(dt.now().year) in x.text]

    # Если последняя дата из БД совпадает с 1-ой найденной на сайте, то возвращаем пустой словарь
    last_db_date_index = industry_dates.index(last_db_date)
    print(f"index: {last_db_date_index}")
    if last_db_date in industry_dates and last_db_date_index == 0:
        print("Последняя дата из БД совпадает с 1-ой найденной на сайте, => словарь пустой")
        return {}
    elif last_db_date_index == 1:
        print(f"    {industry_dates[0]}: ({industry_heads[0]}, {industry_hrefs[0]})")
        return {industry_dates[0]: (industry_heads[0], industry_hrefs[0])}

    heads_hrefs = zip(industry_heads, industry_hrefs)
    industry_dict = dict(zip(industry_dates, heads_hrefs))

    for k in industry_dict:
        if k == last_db_date:
            dict_key = k
    if dict_key:
        del industry_dict[dict_key]

    return industry_dict


def create_file(file_href, year):
    file_name = os.path.split(file_href)[1]
    dir_path = os.path.join(MEDIA, 'industry', f'{year}')
    if not os.path.exists(dir_path):
        os.makedirs(dir_path)

    f_path = os.path.join(dir_path, file_name)

    # Если файла нет, то скачиваем с помощью requests.get
    if not os.path.exists(f_path):
        with open(f_path, 'wb') as f:
            f.write(requests.get(file_href).content)

    # Перевод doc файла в docx
    if not os.path.exists(f"{f_path}x"):
        doc2docx(dir_path)

    docx_file = docx.Document(f_path+'x')

    return docx_file, dir_path


@transaction.atomic
def populate():
    print("-----------------------INDUSTRY BEGIN--------------------------")
    head_dict = last_added_news(HEADER)  # словарь вида dict{date: (title, href)}; 0-title, 1-href
    keys_list = list(reversed(head_dict.keys()))  # сортируем даты по возрастанию
    for k in keys_list:
        title_date = date_int(head_dict[k][0])  # преобразуем "в январе-марте 2024" в список [2024, 1, 3]

        # Заходим в новость, сохраняем файл (запоминаем файл docx и путь к нему)
        docxfile, pwd = create_file(head_dict[k][1], str(title_date[0]))

        # Добавляем в БД, определяем id новости
        news_id = add_news(head_dict[k][0], head_dict[k][1], k)

        # Выкопировка двух таблиц, учитываем январь (в табл.1 визуально меньше столбцов, по факту - нет)
        is_jan = False
        if len(title_date) == 2:
            is_jan = True

        table1, table2 = table_doc(docxfile)

        # Значения из str в float
        table1_float = floating(2, table1)
        table2_float = floating(1, table2)
        create_db(news_id, table1_float, table2_float, is_jan)
        print('--------------------------------------------------------')
        print(f"Добавление данных за {k.strftime('%d.%m.%Y')} завершено")
    print("-----------------------INDUSTRY END--------------------------")

